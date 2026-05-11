"""Convert jobs/<id>/cv.md to jobs/<id>/cv.docx."""

import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SECTION_COLOR = RGBColor(0x1F, 0x49, 0x7D)
META_COLOR = RGBColor(0x60, 0x60, 0x60)


def _add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    hyperlink.append(r)
    paragraph._p.append(hyperlink)


def _parse_inline(paragraph, text, base_size=10, base_color=None):
    """Append inline-formatted runs to paragraph, handling **bold**, *italic*, [link](url)."""
    i = 0
    while i < len(text):
        # Hyperlink
        m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", text[i:])
        if m:
            _add_hyperlink(paragraph, m.group(1), m.group(2))
            i += len(m.group(0))
            continue
        # Bold
        m = re.match(r"\*\*(.+?)\*\*", text[i:])
        if m:
            run = paragraph.add_run(m.group(1))
            run.bold = True
            run.font.size = Pt(base_size)
            i += len(m.group(0))
            continue
        # Italic
        m = re.match(r"\*(.+?)\*", text[i:])
        if m:
            run = paragraph.add_run(m.group(1))
            run.italic = True
            run.font.size = Pt(base_size)
            i += len(m.group(0))
            continue
        # Plain text — consume until next special char
        j = i + 1
        while j < len(text) and text[j] not in ("*", "["):
            j += 1
        run = paragraph.add_run(text[i:j])
        run.font.size = Pt(base_size)
        if base_color:
            run.font.color.rgb = base_color
        i = j


def _add_section_header(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = SECTION_COLOR
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F497D")
    pBdr.append(bottom)
    pPr.append(pBdr)


def render(job_id: str):
    job_dir = Path("jobs") / job_id
    cv_md = job_dir / "cv.md"
    cv_docx = job_dir / "cv.docx"

    if not cv_md.exists():
        print(f"Error: {cv_md} not found. Run /tailor {job_id} first.")
        sys.exit(1)

    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)

    lines = cv_md.read_text(encoding="utf-8").splitlines()
    i = 0
    header_done = False

    while i < len(lines):
        line = lines[i]
        s = line.strip()

        if not s:
            i += 1
            continue

        # Name — first # heading
        if s.startswith("# ") and not s.startswith("## ") and not header_done:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(s[2:].strip())
            run.bold = True
            run.font.size = Pt(18)
            header_done = True
            i += 1
            # Contact line immediately follows
            if i < len(lines) and lines[i].strip() and not lines[i].startswith("#"):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(6)
                _parse_inline(p, lines[i].strip(), base_size=9)
                i += 1
            continue

        # Section header ##
        if s.startswith("## "):
            _add_section_header(doc, s[3:].strip())
            i += 1
            continue

        # Subsection header ###
        if s.startswith("### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(s[4:].strip())
            run.bold = True
            run.font.size = Pt(10)
            i += 1
            continue

        # Italic meta line *text* (dates, locations)
        if s.startswith("*") and s.endswith("*") and not s.startswith("**"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(s[1:-1])
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = META_COLOR
            i += 1
            continue

        # Bullet
        if s.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(1)
            _parse_inline(p, s[2:].strip())
            i += 1
            continue

        # Normal paragraph (skills lines with **Label:** value)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _parse_inline(p, s)
        i += 1

    doc.save(str(cv_docx))
    print(f"Saved: {cv_docx}")


if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python src/render_docx.py <job-id>")
        sys.exit(1)
    render(sys.argv[1])
